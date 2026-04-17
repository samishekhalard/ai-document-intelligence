"""Document ingestion: load -> parse -> chunk -> persist metadata.

Supports PDF, DOCX, TXT/MD, and images (OCR fallback).
"""
from __future__ import annotations

import hashlib
import re
import uuid
from dataclasses import dataclass
from pathlib import Path

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langdetect import DetectorFactory, LangDetectException, detect

from backend.core.config import get_settings
from backend.core.logging import get_logger
from backend.core.models import Chunk, DocumentMetadata, DocumentType, IngestionStatus

DetectorFactory.seed = 42  # deterministic language detection

log = get_logger(__name__)


@dataclass
class LoadedDocument:
    """In-memory representation of a parsed document."""

    document_id: str
    title: str
    source_path: str
    file_type: str
    text: str
    size_bytes: int


_SUPPORTED_EXT = {".pdf", ".docx", ".txt", ".md", ".png", ".jpg", ".jpeg", ".tif", ".tiff"}


def _document_id(path: Path) -> str:
    """Stable document ID derived from absolute path + size."""
    h = hashlib.sha1()
    h.update(str(path.resolve()).encode("utf-8"))
    h.update(str(path.stat().st_size).encode("utf-8"))
    return h.hexdigest()[:16]


def _load_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # pragma: no cover - defensive
            log.warning(f"pdf page {i} failed: {exc}")
    text = "\n\n".join(pages).strip()
    if text:
        return text
    # Empty PDF text -> try OCR via pdf2image + tesseract
    log.info(f"PDF {path.name} has no text layer, attempting OCR")
    from backend.pipelines.ocr_pipeline import ocr_pdf

    return ocr_pdf(path)


def _load_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text)


def _load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _load_image(path: Path) -> str:
    from backend.pipelines.ocr_pipeline import ocr_image

    return ocr_image(path)


def load_document(path: Path) -> LoadedDocument:
    """Parse a single document from disk into a :class:`LoadedDocument`."""
    if not path.exists():
        raise FileNotFoundError(path)
    ext = path.suffix.lower()
    if ext not in _SUPPORTED_EXT:
        raise ValueError(f"Unsupported file type: {ext}")

    if ext == ".pdf":
        text = _load_pdf(path)
    elif ext == ".docx":
        text = _load_docx(path)
    elif ext in {".txt", ".md"}:
        text = _load_text(path)
    else:
        text = _load_image(path)

    text = _normalise_whitespace(text)
    if not text.strip():
        raise ValueError(f"No extractable text in {path}")

    return LoadedDocument(
        document_id=_document_id(path),
        title=path.stem.replace("_", " ").title(),
        source_path=str(path.resolve()),
        file_type=ext.lstrip("."),
        text=text,
        size_bytes=path.stat().st_size,
    )


def _normalise_whitespace(text: str) -> str:
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def detect_language(text: str) -> str:
    """Return ISO language code (``en``, ``ru``, ...) or ``'unknown'``."""
    try:
        return detect(text[:2000])
    except LangDetectException:
        return "unknown"


def _classify_heuristic(text: str) -> DocumentType:
    """Fast keyword-based document classifier used during ingestion."""
    t = text.lower()
    scores = {
        DocumentType.LEGAL: sum(
            kw in t for kw in ("agreement", "party", "whereas", "hereby",
                                 "clause", "jurisdiction", "договор", "стороны")
        ),
        DocumentType.FINANCIAL: sum(
            kw in t for kw in ("revenue", "profit", "balance", "fiscal",
                                 "ebitda", "usd", "eur", "выручка", "прибыль",
                                 "баланс", "отчет")
        ),
        DocumentType.TECHNICAL: sum(
            kw in t for kw in ("api", "system", "architecture", "module",
                                 "specification", "latency", "throughput",
                                 "database")
        ),
    }
    best = max(scores.items(), key=lambda kv: kv[1])
    return best[0] if best[1] >= 2 else DocumentType.GENERAL


def chunk_text(text: str, document_id: str) -> list[Chunk]:
    """Split ``text`` using recursive character splitter with settings sizes."""
    settings = get_settings()
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
        length_function=len,
    )
    pieces = splitter.split_text(text)
    chunks: list[Chunk] = []
    for idx, piece in enumerate(pieces):
        if not piece.strip():
            continue
        chunks.append(
            Chunk(
                chunk_id=f"{document_id}:{idx:04d}:{uuid.uuid4().hex[:6]}",
                document_id=document_id,
                text=piece,
                index=idx,
                metadata={"char_len": len(piece)},
            )
        )
    return chunks


class IngestionPipeline:
    """Top-level orchestrator: file -> chunks + metadata."""

    def __init__(self) -> None:
        self.settings = get_settings()

    def prepare(self, path: Path) -> tuple[DocumentMetadata, list[Chunk]]:
        """Parse, chunk, and characterise a document on disk.

        Runs the full pre-embedding pipeline: format-specific extraction,
        whitespace normalisation, recursive character splitting, language
        detection, and keyword-based classification.

        Args:
            path: Absolute or relative filesystem path to a supported file.

        Returns:
            A tuple of ``(DocumentMetadata, list[Chunk])`` ready to be
            handed to :meth:`EmbeddingPipeline.add_chunks`.

        Raises:
            FileNotFoundError: If ``path`` does not exist.
            ValueError: If the file extension is unsupported or the
                document yields no extractable text.
        """
        log.info(f"Ingesting {path.name}")
        loaded = load_document(path)
        language = detect_language(loaded.text)
        doc_type = _classify_heuristic(loaded.text)

        chunks = chunk_text(loaded.text, loaded.document_id)
        for c in chunks:
            c.metadata.update(
                {
                    "document_id": loaded.document_id,
                    "title": loaded.title,
                    "file_type": loaded.file_type,
                    "language": language,
                    "doc_type": doc_type.value,
                    "source_path": loaded.source_path,
                }
            )

        meta = DocumentMetadata(
            document_id=loaded.document_id,
            title=loaded.title,
            source_path=loaded.source_path,
            file_type=loaded.file_type,
            language=language,
            doc_type=doc_type,
            n_chunks=len(chunks),
            status=IngestionStatus.PROCESSING,
            size_bytes=loaded.size_bytes,
        )
        log.info(
            f"Prepared {len(chunks)} chunks for '{loaded.title}' "
            f"(lang={language}, type={doc_type.value})"
        )
        return meta, chunks
